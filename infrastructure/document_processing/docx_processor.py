"""
DOCX Document Processor.

Concrete implementation of BaseDocumentProcessor for Word documents.
Uses python-docx for structured extraction (paragraphs, tables)
with docx2txt as a fast fallback.

Design decisions:
- python-docx gives access to document structure (headings, tables, lists)
- Extract tables as pipe-delimited text to preserve tabular data
- docx2txt fallback handles malformed .docx that python-docx rejects
"""

from __future__ import annotations

from pathlib import Path

from config.logging_config import get_logger
from core.interfaces.document_processor import (
    BaseDocumentProcessor,
    DocumentProcessingError,
    ExtractedDocument,
)
from infrastructure.document_processing.text_cleaner import (
    clean_extracted_text,
    count_words,
    truncate_text,
)

logger = get_logger(__name__)


class DOCXProcessor(BaseDocumentProcessor):
    """
    DOCX text extractor using python-docx with docx2txt fallback.

    Extracts paragraphs and tables in document order, producing
    a clean linear text representation suitable for NLP processing.
    """

    SUPPORTED_EXTENSION = "docx"
    MAX_CHARS = 100_000

    def supports(self, file_extension: str) -> bool:
        """Return True for 'docx' extension."""
        return file_extension.lower().strip(".") == self.SUPPORTED_EXTENSION

    def extract_text(self, file_path: str) -> ExtractedDocument:
        """
        Extract and clean text from a DOCX file.

        Args:
            file_path: Absolute or relative path to the DOCX.

        Returns:
            ExtractedDocument with cleaned text.

        Raises:
            DocumentProcessingError: If the file cannot be read.
        """
        path = Path(file_path)
        self._validate_file(path)

        logger.info("Extracting DOCX", extra={"file": path.name})

        # Try python-docx first for structured extraction
        raw_text, method = self._extract_with_python_docx(path)

        # Fall back to docx2txt if python-docx returned nothing
        if not raw_text.strip():
            logger.warning(
                "python-docx returned empty text — trying docx2txt fallback",
                extra={"file": path.name},
            )
            raw_text, method = self._extract_with_docx2txt(path)

        if not raw_text.strip():
            raise DocumentProcessingError(
                f"No text could be extracted from '{path.name}'."
            )

        cleaned = clean_extracted_text(raw_text)
        cleaned = truncate_text(cleaned, self.MAX_CHARS)
        word_count = count_words(cleaned)

        logger.info(
            "DOCX extraction complete",
            extra={"file": path.name, "words": word_count, "method": method},
        )

        return ExtractedDocument(
            file_name=path.name,
            file_path=str(path.resolve()),
            raw_text=cleaned,
            page_count=0,  # DOCX doesn't have fixed page count at extraction
            word_count=word_count,
            extraction_method=method,
        )

    def _extract_with_python_docx(self, path: Path) -> tuple[str, str]:
        """
        Extract text using python-docx, preserving paragraph and table structure.

        Returns:
            Tuple of (extracted_text, method_name).
        """
        try:
            from docx import Document  # python-docx

            doc = Document(str(path))
            sections: list[str] = []

            for block in self._iter_block_items(doc):
                if hasattr(block, "text"):
                    # It's a paragraph
                    text = block.text.strip()
                    if text:
                        sections.append(text)
                else:
                    # It's a table — render as pipe-delimited rows
                    table_lines = self._extract_table(block)
                    if table_lines:
                        sections.extend(table_lines)

            return "\n".join(sections), "python-docx"

        except Exception as e:
            logger.warning(
                "python-docx extraction failed",
                extra={"error": str(e)},
            )
            return "", "python-docx-failed"

    def _extract_with_docx2txt(self, path: Path) -> tuple[str, str]:
        """
        Fallback extraction using docx2txt.

        Returns:
            Tuple of (extracted_text, method_name).
        """
        try:
            import docx2txt

            text = docx2txt.process(str(path))
            return text or "", "docx2txt"
        except Exception as e:
            logger.error("docx2txt extraction failed", extra={"error": str(e)})
            return "", "docx2txt-failed"

    def _iter_block_items(self, doc: object):
        """
        Yield paragraphs and tables in document order.

        python-docx's Document.paragraphs and Document.tables are
        separate lists — this interleaves them in their original order
        by walking the XML body directly.
        """
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        body = doc.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                yield Table(child, doc)

    def _extract_table(self, table: object) -> list[str]:
        """
        Render a DOCX table as pipe-delimited text lines.

        Args:
            table: python-docx Table object.

        Returns:
            List of pipe-delimited row strings.
        """
        lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Deduplicate merged cells (python-docx repeats content for merged cells)
            unique_cells = []
            for cell in cells:
                if not unique_cells or cell != unique_cells[-1]:
                    unique_cells.append(cell)
            line = " | ".join(c for c in unique_cells if c)
            if line:
                lines.append(line)
        return lines

    def _validate_file(self, path: Path) -> None:
        """Validate that the file exists and is a DOCX."""
        if not path.exists():
            raise DocumentProcessingError(f"File not found: '{path}'")
        if not path.is_file():
            raise DocumentProcessingError(f"Path is not a file: '{path}'")
        if path.suffix.lower() != ".docx":
            raise DocumentProcessingError(
                f"Expected a .docx file, got: '{path.suffix}'"
            )
        if path.stat().st_size == 0:
            raise DocumentProcessingError(f"File is empty: '{path.name}'")
